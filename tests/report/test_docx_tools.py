from io import BytesIO

import pytest
from docx import Document

from policybot.report.docx_tools import (
    add_citation_paragraph,
    add_hyperlink,
    replace_paragraph_text,
    table_by_heading,
    tick_checkbox,
)


def test_table_by_heading_is_not_position_dependent():
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Intrus"
    expected = document.add_table(rows=1, cols=1)
    expected.cell(0, 0).text = "Évaluation — critères"

    assert table_by_heading(document, "evaluation - criteres")._tbl is expected._tbl


def test_table_by_heading_fails_loudly():
    with pytest.raises(RuntimeError, match="Absent"):
        table_by_heading(Document(), "Absent")


def test_tick_checkbox_handles_all_template_levels():
    document = Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "☐ F"
    for letter in "MEC":
        cell.add_paragraph(f"☐ {letter}")

    tick_checkbox(cell.paragraphs, "C")

    assert [paragraph.text for paragraph in cell.paragraphs] == [
        "☐ F", "☐ M", "☐ E", "☒ C",
    ]


def test_replace_paragraph_text_operates_across_runs():
    paragraph = Document().add_paragraph()
    paragraph.add_run("Rens. per")
    paragraph.add_run("sonnels : ☐ Oui")

    replace_paragraph_text(paragraph, "☐ Oui", "☒ Oui")

    assert paragraph.text == "Rens. personnels : ☒ Oui"
    assert paragraph.runs[1].text == ""


def test_add_citation_paragraph_creates_an_external_hyperlink():
    document = Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    add_citation_paragraph(
        cell, "Officielle", "Citation exacte", "https://vendor.test/source"
    )
    buffer = BytesIO()
    document.save(buffer)
    reopened = Document(buffer)

    assert "Officielle — « Citation exacte »" in reopened.tables[0].cell(0, 0).text
    assert "vendor.test/source" in reopened.tables[0].cell(0, 0).text
    assert any(
        relationship.is_external
        and relationship.target_ref == "https://vendor.test/source"
        for relationship in reopened.part.rels.values()
    )


def test_add_hyperlink_can_use_a_short_display_label():
    document = Document()
    paragraph = document.add_paragraph()
    add_hyperlink(
        paragraph,
        "https://vendor.test/source#:~:text=long-fragment",
        "Ouvrir la source",
        bold=True,
    )
    buffer = BytesIO()
    document.save(buffer)
    reopened = Document(buffer)

    assert reopened.paragraphs[0].text == "Ouvrir la source"
    assert "long-fragment" not in reopened.paragraphs[0].text
    assert any(
        relationship.target_ref.endswith("long-fragment")
        for relationship in reopened.part.rels.values()
        if relationship.is_external
    )
