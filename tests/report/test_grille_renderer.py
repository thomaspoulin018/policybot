import copy
from datetime import date
from io import BytesIO
import zipfile

from docx import Document
from docx.oxml.ns import qn

from policybot.contract.criteres import CRITERIA
from policybot.models import (
    CriterionCitation,
    CriterionFinding,
    InterviewState,
    RequestInfo,
    ToolRef,
    Usage,
)
from policybot.report.docx_tools import _normalized_text
from policybot.report.grille import (
    _RISK_COLUMN_WIDTHS,
    _SOURCE_COLUMN_WIDTHS,
    grille_template_path,
    render_grille,
)


def _finding(partie: str, risk="M", *, outcome="ok") -> CriterionFinding:
    definition = next(item for item in CRITERIA if item.partie == partie)
    return CriterionFinding(
        id=definition.id,
        partie=partie,
        category=definition.category,
        criterion=definition.criterion,
        question=definition.question,
        answer=f"Réponse factuelle {partie}.",
        inherent_risk=risk,
        justification=f"Justification {partie}.",
        outcome=outcome,
        rejected_citations=1,
        cost_dollars=0.02,
        citations=[
            CriterionCitation(
                url=f"https://vendor.test/{partie.lower()}",
                deep_link=f"https://vendor.test/{partie.lower()}#preuve",
                title=f"Contrat {partie}",
                text=f"Citation {partie}",
                source_type="official",
                anchored=True,
            )
        ],
    )


def _state(usage_count=1) -> InterviewState:
    usages = [
        Usage(
            description=f"Usage {index + 1}",
            data_classification=("Protégé A" if index % 2 == 0 else "Non classifié"),
            rens_personnels=index % 2 == 0,
        )
        for index in range(usage_count)
    ]
    return InterviewState(
        interview_id="grille",
        status="complete",
        request=RequestInfo(numero="REQ-42", date=date(2026, 8, 12)),
        tools=[ToolRef(name="Outil X", findings=[_finding("A"), _finding("B", "E")])],
        usages=usages,
    )


def _criterion_row(table, label: str):
    wanted = _normalized_text(label)
    return next(row for row in table.rows if _normalized_text(row.cells[0].text) == wanted)


def _tables_starting(document, heading: str):
    wanted = _normalized_text(heading)
    return [
        table for table in document.tables
        if table.rows and _normalized_text(table.cell(0, 0).text).startswith(wanted)
    ]


def test_grille_fills_identification_risk_justification_and_sources():
    payload = render_grille(_state())
    assert payload.startswith(b"PK")
    document = Document(BytesIO(payload))

    identification = document.tables[0].rows[0].cells
    assert identification[0].paragraphs[1].text == "REQ-42"
    assert identification[2].paragraphs[1].text == "Outil X"
    assert identification[4].paragraphs[1].text == "12 août 2026"

    definition = next(item for item in CRITERIA if item.partie == "A")
    row = _criterion_row(_tables_starting(document, "Critère")[0], definition.criterion)
    assert "☒ M" in row.cells[2].text
    assert "Justification : Justification A." in row.cells[2].text
    assert "Justification A." in row.cells[2].text
    assert "Réponse factuelle A." in row.cells[6].text
    assert "Officielle — « Citation A »" in row.cells[6].text
    assert "Justification A." not in row.cells[6].text
    assert "1 citation(s) non ancrée(s) rejetée(s)." in row.cells[6].text
    assert any(rel.is_external for rel in document.part.rels.values())


def test_generated_risk_tables_use_fixed_balanced_geometry():
    document = Document(BytesIO(render_grille(_state())))
    for table in (
        _tables_starting(document, "Critère")[0],
        _tables_starting(document, "Catégorie")[0],
    ):
        assert table.autofit is False
        assert [
            int(column.get(qn("w:w")))
            for column in table._tbl.tblGrid.gridCol_lst
        ] == list(_RISK_COLUMN_WIDTHS)
        assert table.cell(2, 2).width.twips == _RISK_COLUMN_WIDTHS[2]
        assert table.cell(2, 6).width.twips == _RISK_COLUMN_WIDTHS[6]


def test_source_register_is_a_styled_fixed_width_table():
    document = Document(BytesIO(render_grille(_state())))
    table = next(table for table in document.tables if table.cell(0, 0).text == "TYPE")

    assert table.autofit is False
    assert [
        int(column.get(qn("w:w")))
        for column in table._tbl.tblGrid.gridCol_lst
    ] == list(_SOURCE_COLUMN_WIDTHS)
    assert table.rows[0]._tr.xpath("./w:trPr/w:tblHeader")
    assert table.cell(0, 0)._tc.xpath("./w:tcPr/w:shd[@w:fill='1F4E78']")
    assert table.cell(1, 3).text.startswith("Ouvrir la source")
    assert "#:~:text=" not in table.cell(1, 3).text
    assert any("référence(s) unique(s)" in paragraph.text for paragraph in document.paragraphs)


def test_one_part_b_block_is_cloned_and_filled_per_usage():
    document = Document(BytesIO(render_grille(_state(3))))
    usage_tables = _tables_starting(document, "Catégorie")
    headings = [paragraph.text for paragraph in document.paragraphs if "Usage évalué" in paragraph.text]

    assert len(usage_tables) == 3
    assert len(headings) == 3
    assert all(f"Usage {index}" in headings[index - 1] for index in range(1, 4))
    assert "☒ Protégé A" in headings[0] and "☒ Oui" in headings[0]
    assert "☒ Non classifié" in headings[1] and "☒ Non" in headings[1]
    for table in usage_tables:
        definition = next(item for item in CRITERIA if item.partie == "B")
        row = _criterion_row(table, definition.criterion)
        assert "☒ E" in row.cells[2].text


def test_cloning_does_not_duplicate_comment_identifiers():
    payload = render_grille(_state(4))
    with zipfile.ZipFile(BytesIO(payload)) as archive:
        document_xml = archive.read("word/document.xml")

    assert document_xml.count(b"commentRangeStart") == 4
    assert document_xml.count(b"commentRangeEnd") == 4
    assert document_xml.count(b"commentReference") == 4


def test_zero_usage_leaves_the_official_usage_block_unfilled():
    document = Document(BytesIO(render_grille(_state(0))))
    headings = [paragraph.text for paragraph in document.paragraphs if "Usage évalué" in paragraph.text]

    assert len(_tables_starting(document, "Catégorie")) == 1
    assert len(headings) == 1
    assert "_______________________________________" in headings[0]
    assert "☒" not in headings[0]


def test_part_c_and_prefilled_mitigation_cells_are_preserved():
    template = Document(grille_template_path())
    rendered = Document(BytesIO(render_grille(_state())))

    assert rendered.tables[5]._tbl.xml == template.tables[5]._tbl.xml
    usage_table = _tables_starting(rendered, "Catégorie")[0]
    assert "valider l’exactitude" in usage_table.cell(7, 3).text
    assert "minimiser les risques de biais" in usage_table.cell(8, 3).text


def test_grille_filling_survives_a_table_inserted_before_the_template(monkeypatch, tmp_path):
    document = Document(grille_template_path())
    intruder = copy.deepcopy(document.tables[0]._tbl)
    for text_node in intruder.iter(qn("w:t")):
        text_node.text = "Tableau ajouté après coup"
    document._body._element.insert(0, intruder)
    path = tmp_path / "grille-avec-intrus.docx"
    document.save(path)
    monkeypatch.setattr("policybot.report.grille.grille_template_path", lambda: path)

    rendered = Document(BytesIO(render_grille(_state())))

    assert rendered.tables[0].cell(0, 0).text.startswith("Tableau ajouté")
    assert rendered.tables[1].cell(0, 0).paragraphs[1].text == "REQ-42"


def test_missing_criterion_fails_loudly(monkeypatch, tmp_path):
    document = Document(grille_template_path())
    document.tables[2].cell(2, 0).text = "Critère retiré"
    path = tmp_path / "grille-amputee.docx"
    document.save(path)
    monkeypatch.setattr("policybot.report.grille.grille_template_path", lambda: path)

    try:
        render_grille(_state())
    except RuntimeError as error:
        assert "absent du gabarit" in str(error)
    else:
        raise AssertionError("Le critère manquant aurait dû interrompre le rendu.")
