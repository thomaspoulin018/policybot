import copy
import xml.etree.ElementTree as ET
import zipfile
from io import BytesIO

import pytest
from docx import Document

from policybot.contract.criteres import CRITERIA
from policybot.models import (
    CriterionCitation, CriterionFinding, InterviewState, RequestInfo, ToolRef,
)
from policybot.report.renderer import (
    _WORD_NS, _ordered_rows, fiche_template_path, render_docx,
)


def _state():
    definition = next(item for item in CRITERIA if item.partie == "A")
    finding = CriterionFinding(
        id=definition.id, partie="A", category=definition.category,
        criterion=definition.criterion, question=definition.question,
        answer="Réponse vérifiée.", inherent_risk="M",
        justification="Justification.", cost_dollars=0.02,
        citations=[CriterionCitation(
            url="https://vendor.test/page", text="Citation exacte",
            anchored=True,
            deep_link="https://vendor.test/page#:~:text=Citation%20exacte",
        )],
    )
    return InterviewState(
        interview_id="i", status="complete", request=RequestInfo(numero="REQ-1"),
        tools=[ToolRef(name="ToolX", iag_type="publique", findings=[finding])],
    )


def test_question_posee_et_question_rendue_proviennent_du_meme_yaml():
    definition = next(item for item in CRITERIA if item.exa is not None)
    finding = CriterionFinding(
        id=definition.id,
        partie=definition.partie,
        category="Ancienne catégorie qui ne correspond plus",
        criterion="Ancien nom qui ne correspond plus",
        question=definition.question,
    )

    row = _ordered_rows([finding], [definition])[0]

    assert row["finding"] is finding
    assert row["description"] == finding.question == definition.question


def test_fiche_export_uses_the_official_docx_template():
    docx = render_docx(_state())
    assert docx.startswith(b"PK")
    document = Document(BytesIO(docx))
    assert len(document.tables) == 8
    assert [len(table.rows) for table in document.tables] == [7, 8, 21, 5, 6, 5, 5, 5]
    assert document.tables[0].cell(0, 1).text == "REQ-1"
    assert document.tables[1].cell(0, 1).text == "ToolX"
    # La recommandation demeure réservée à l'autorité désignée.
    assert document.tables[7].cell(3, 1).text == ""


def test_section_8_remains_empty_even_when_findings_exist():
    document = Document(BytesIO(render_docx(_state())))
    section_8 = document.tables[7]

    assert len(section_8.rows) == 5
    assert [section_8.cell(row, 1).text for row in range(5)] == [""] * 5


def _template_with_an_extra_table_in_front(original: bytes) -> bytes:
    """Insère un tableau supplémentaire avant les autres, comme le ferait un
    éditeur qui enrichit le gabarit institutionnel."""
    root = ET.fromstring(original)
    body = root.find(f"{{{_WORD_NS}}}body")
    first = body.find(f"{{{_WORD_NS}}}tbl")
    intruder = copy.deepcopy(first)
    for node in intruder.iter(f"{{{_WORD_NS}}}t"):
        node.text = "Tableau ajouté après coup"
    body.insert(list(body).index(first), intruder)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def test_docx_filling_survives_a_table_inserted_into_the_template(monkeypatch, tmp_path):
    source = fiche_template_path()
    patched = tmp_path / "gabarit.docx"
    with zipfile.ZipFile(source) as original:
        with zipfile.ZipFile(patched, "w", zipfile.ZIP_DEFLATED) as target:
            for item in original.infolist():
                payload = original.read(item.filename)
                if item.filename == "word/document.xml":
                    payload = _template_with_an_extra_table_in_front(payload)
                target.writestr(item, payload)

    monkeypatch.setattr(
        "policybot.report.renderer.fiche_template_path", lambda: patched,
    )
    document = Document(BytesIO(render_docx(_state())))

    # Le tableau intrus reste vide et les valeurs vont toujours au bon endroit.
    assert len(document.tables) == 9
    assert document.tables[0].cell(0, 1).text.startswith("Tableau ajouté")
    assert document.tables[1].cell(0, 1).text == "REQ-1"
    assert document.tables[2].cell(0, 1).text == "ToolX"


def test_a_template_missing_a_table_fails_loudly(monkeypatch, tmp_path):
    source = fiche_template_path()
    patched = tmp_path / "ampute.docx"
    with zipfile.ZipFile(source) as original:
        with zipfile.ZipFile(patched, "w", zipfile.ZIP_DEFLATED) as target:
            for item in original.infolist():
                payload = original.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(payload)
                    body = root.find(f"{{{_WORD_NS}}}body")
                    body.remove(body.find(f"{{{_WORD_NS}}}tbl"))
                    payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                target.writestr(item, payload)

    monkeypatch.setattr(
        "policybot.report.renderer.fiche_template_path", lambda: patched,
    )
    with pytest.raises(RuntimeError, match="Numéro de demande"):
        render_docx(_state())
