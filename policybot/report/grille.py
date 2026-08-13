from __future__ import annotations

import copy
from io import BytesIO
import os
from pathlib import Path
import re
from urllib.parse import urlsplit

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from policybot.contract.criteres import CRITERIA
from policybot.models import CriterionFinding, InterviewState
from policybot.report.docx_tools import (
    _normalized_text,
    add_citation_paragraph,
    add_hyperlink,
    replace_paragraph_text,
    table_by_heading,
    tick_checkbox,
)
from policybot.report.renderer import (
    _filename_stem,
    _format_french_date,
    _source_type_label,
    docx_output_dir,
)


_DEFAULT_GRILLE_TEMPLATE = (
    Path(__file__).resolve().parents[2]
    / "documents_reference"
    / "SI_-_Grille_valuation_des_risques.docx"
)
_CLONE_MARKERS = (
    "bookmarkStart",
    "bookmarkEnd",
    "commentRangeStart",
    "commentRangeEnd",
    "commentReference",
)
# Le gabarit consacre seulement 799 DXA (0,55 po) au risque inhérent, alors
# que cette cellule reçoit désormais une justification. On conserve la largeur
# totale officielle de 14 040 DXA et on rééquilibre uniquement les colonnes 2
# et 6, qui portent le contenu généré par PolicyBot.
_RISK_COLUMN_WIDTHS = (1610, 2730, 1900, 2175, 799, 1321, 3505)
_SOURCE_COLUMN_WIDTHS = (1400, 3650, 6890, 2100)
_SOURCE_HEADER_FILL = "1F4E78"
_SOURCE_ALT_FILL = "F3F7FA"
_SOURCE_BORDER = "B8C7D1"


def grille_template_path() -> Path:
    return Path(os.environ.get("POLICYBOT_GRILLE_TEMPLATE") or _DEFAULT_GRILLE_TEMPLATE)


def grille_filename(state: InterviewState) -> str:
    return f"{_filename_stem(state)}-grille.docx"


def _first_tool(state: InterviewState):
    return state.tools[0] if state.tools else None


def _findings(state: InterviewState, partie: str) -> list[CriterionFinding]:
    tool = _first_tool(state)
    return [item for item in (tool.findings if tool else []) if item.partie == partie]


def _set_identification_value(table, heading: str, value: object | None) -> None:
    wanted = _normalized_text(heading)
    for cell in table.rows[0].cells:
        if cell.paragraphs and _normalized_text(cell.paragraphs[0].text).startswith(wanted):
            paragraph = cell.paragraphs[1] if len(cell.paragraphs) > 1 else cell.add_paragraph()
            paragraph.text = "" if value is None else str(value)
            return
    raise RuntimeError(f"Champ d'identification introuvable : « {heading} ».")


def _fill_identification(document, state: InterviewState) -> None:
    table = table_by_heading(document, "Numéro demande")
    tool = _first_tool(state)
    _set_identification_value(table, "Numéro demande", state.request.numero)
    _set_identification_value(table, "Outil évalué", tool.name if tool else "")
    _set_identification_value(table, "Date", _format_french_date(state.request.date))


def _row_index(table) -> dict[str, int]:
    return {
        _normalized_text(row.cells[0].text): index
        for index, row in enumerate(table.rows)
        if row.cells and row.cells[0].text.strip()
    }


def _clear_cell(cell) -> None:
    cell.text = ""


def _set_width(element, width: int) -> None:
    element.set(qn("w:w"), str(width))
    element.set(qn("w:type"), "dxa")


def _apply_table_geometry(table, widths: tuple[int, ...]) -> None:
    """Apply fixed, internally consistent DXA widths to a Word table."""
    table.autofit = False
    table_width = sum(widths)

    table_properties = table._tbl.tblPr
    table_width_node = next(
        (child for child in table_properties if child.tag == qn("w:tblW")), None
    )
    if table_width_node is None:
        table_width_node = table_properties._add_tblW()
    _set_width(table_width_node, table_width)
    layout = next(
        (child for child in table_properties if child.tag == qn("w:tblLayout")), None
    )
    if layout is None:
        layout = table_properties._add_tblLayout()
    layout.set(qn("w:type"), "fixed")

    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths, strict=True):
        grid_column.set(qn("w:w"), str(width))

    for row in table._tbl.tr_lst:
        column_index = 0
        for cell in row.tc_lst:
            span_node = cell.tcPr.gridSpan
            span = int(span_node.val) if span_node is not None else 1
            cell_width = sum(widths[column_index:column_index + span])
            width_node = cell.tcPr.tcW
            if width_node is None:
                width_node = cell.tcPr._add_tcW()
            _set_width(width_node, cell_width)
            column_index += span


def _apply_risk_table_geometry(table) -> None:
    _apply_table_geometry(table, _RISK_COLUMN_WIDTHS)


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = next(
        (child for child in properties if child.tag == qn("w:shd")), None
    )
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def _set_cell_margins(cell, *, top=100, start=110, bottom=100, end=110) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = next(
        (child for child in properties if child.tag == qn("w:tcMar")), None
    )
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = next(
            (child for child in margins if child.tag == qn(f"w:{tag}")), None
        )
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = next(
        (child for child in properties if child.tag == qn("w:tblBorders")), None
    )
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for tag in ("top", "start", "bottom", "end", "insideH", "insideV"):
        border = next(
            (child for child in borders if child.tag == qn(f"w:{tag}")), None
        )
        if border is None:
            border = OxmlElement(f"w:{tag}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), _SOURCE_BORDER)


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if not any(child.tag == qn("w:tblHeader") for child in properties):
        properties.append(OxmlElement("w:tblHeader"))


def _style_source_paragraph(paragraph, size=8.25, *, bold=False, color=None) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _compact_excerpt(value: str, limit: int = 360) -> str:
    text = " ".join(value.split())
    if len(text) <= limit:
        return text
    return text[: limit - 1].rsplit(" ", 1)[0] + "…"


def _format_generated_paragraph(paragraph, size: float = 8.5) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.size = Pt(size)


def _fill_finding_row(table, row_index: int, finding: CriterionFinding) -> None:
    risk_cell = table.cell(row_index, 2)
    if finding.inherent_risk:
        tick_checkbox(risk_cell.paragraphs, finding.inherent_risk)
    for checkbox in risk_cell.paragraphs[:4]:
        _format_generated_paragraph(checkbox, 8)
    if finding.justification.strip():
        justification = risk_cell.add_paragraph()
        justification.paragraph_format.space_before = Pt(4)
        label = justification.add_run("Justification : ")
        label.bold = True
        justification.add_run(finding.justification.strip())
        _format_generated_paragraph(justification, 8)
        justification.paragraph_format.space_before = Pt(4)

    observation_cell = table.cell(row_index, 6)
    _clear_cell(observation_cell)
    answer = (
        finding.answer.strip()
        if finding.outcome == "ok" and finding.answer.strip()
        else "Aucune source probante trouvée."
    )
    observation_cell.paragraphs[0].text = answer
    _format_generated_paragraph(observation_cell.paragraphs[0], 8.5)
    for citation in finding.citations:
        add_citation_paragraph(
            observation_cell,
            _source_type_label(citation.source_type),
            citation.text,
            citation.deep_link or citation.url,
        )
    if finding.rejected_citations:
        rejected = observation_cell.add_paragraph(
            f"{finding.rejected_citations} citation(s) non ancrée(s) rejetée(s)."
        )
        _format_generated_paragraph(rejected, 8)
        for run in rejected.runs:
            run.italic = True


def _fill_risk_table(table, state: InterviewState, partie: str) -> None:
    _apply_risk_table_geometry(table)
    rows = _row_index(table)
    definitions = [item for item in CRITERIA if item.partie == partie]
    finding_by_id = {item.id: item for item in _findings(state, partie)}
    for definition in definitions:
        key = _normalized_text(definition.criterion)
        if key not in rows:
            raise RuntimeError(
                f"Le critère « {definition.criterion} » ({definition.id}) est absent "
                "du gabarit Word de grille d'évaluation."
            )
        finding = finding_by_id.get(definition.id)
        if finding is not None:
            _fill_finding_row(table, rows[key], finding)


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _scrub_clone_identifiers(element) -> None:
    for marker in _CLONE_MARKERS:
        for node in list(element.iter(qn(f"w:{marker}"))):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)


def _usage_blocks(document, count: int) -> list[tuple[Paragraph, Table]]:
    body = document._body._element
    base_paragraph = next(
        (
            child for child in body.iterchildren()
            if child.tag == qn("w:p")
            and _normalized_text(_element_text(child)).startswith("usage evalue")
        ),
        None,
    )
    if base_paragraph is None or base_paragraph.getnext() is None:
        raise RuntimeError("Bloc « Usage évalué » introuvable dans le gabarit de grille.")
    base_table = base_paragraph.getnext()
    if base_table.tag != qn("w:tbl"):
        raise RuntimeError("Le tableau Partie B ne suit pas le libellé « Usage évalué ».")

    blocks = [(Paragraph(base_paragraph, document._body), Table(base_table, document._body))]
    cursor = base_table
    for _ in range(max(0, count - 1)):
        paragraph_copy = copy.deepcopy(base_paragraph)
        table_copy = copy.deepcopy(base_table)
        _scrub_clone_identifiers(paragraph_copy)
        _scrub_clone_identifiers(table_copy)
        cursor.addnext(paragraph_copy)
        paragraph_copy.addnext(table_copy)
        cursor = table_copy
        blocks.append(
            (Paragraph(paragraph_copy, document._body), Table(table_copy, document._body))
        )
    return blocks


def _fill_usage_heading(paragraph, usage) -> None:
    placeholder = re.search(r"_{3,}", paragraph.text)
    if not placeholder:
        raise RuntimeError("Zone de description introuvable dans « Usage évalué ».")
    replace_paragraph_text(paragraph, placeholder.group(0), usage.description or "")
    if usage.data_classification:
        replace_paragraph_text(
            paragraph,
            f"☐ {usage.data_classification}",
            f"☒ {usage.data_classification}",
        )
    personal_label = "Oui" if usage.rens_personnels else "Non"
    replace_paragraph_text(paragraph, f"☐ {personal_label}", f"☒ {personal_label}")


def _append_source_register(document, state: InterviewState) -> None:
    tool = _first_tool(state)
    findings = tool.findings if tool else []
    if not findings:
        return
    seen: set[str] = set()
    sources = []
    for finding in findings:
        for citation in finding.citations:
            if citation.url and citation.url not in seen:
                seen.add(citation.url)
                sources.append(citation)

    sources.sort(
        key=lambda item: (
            0 if item.source_type == "official" else 1,
            (item.title or item.url).casefold(),
        )
    )

    heading = document.add_heading("Sources contractuelles consultées", level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(3)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(_SOURCE_HEADER_FILL)

    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(7)
    summary_run = summary.add_run(
        f"{len(sources)} référence(s) unique(s)  •  "
        f"Liens actifs  •  Coût Exa total : {tool.total_cost_dollars:.2f} $"
    )
    summary_run.font.size = Pt(8.5)
    summary_run.font.color.rgb = RGBColor.from_string("536878")

    table = document.add_table(rows=1, cols=4)
    for cell, text in zip(
        table.rows[0].cells,
        ("TYPE", "SOURCE", "EXTRAIT PROBANT", "ACCÈS"),
        strict=True,
    ):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, _SOURCE_HEADER_FILL)
        _set_cell_margins(cell, top=120, bottom=120)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_source_paragraph(paragraph, 8.5, bold=True, color="FFFFFF")
    _repeat_table_header(table.rows[0])

    type_fills = {
        "official": "DDEBF7",
        "other": "FFF2CC",
        "unknown": "E7E6E6",
    }
    if sources:
        for index, citation in enumerate(sources):
            cells = table.add_row().cells
            for cell in cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                _set_cell_margins(cell)
                if index % 2:
                    _set_cell_shading(cell, _SOURCE_ALT_FILL)

            cells[0].text = _source_type_label(citation.source_type).upper()
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cells[0], type_fills.get(citation.source_type, "E7E6E6"))
            _style_source_paragraph(cells[0].paragraphs[0], 7.5, bold=True, color="294861")

            cells[1].text = citation.title or urlsplit(citation.url).netloc
            _style_source_paragraph(cells[1].paragraphs[0], 8.5, bold=True, color="1F4E78")
            date_line = cells[1].add_paragraph(
                f"Consultée le {_format_french_date(citation.collected_at)}"
            )
            date_line.paragraph_format.space_before = Pt(2)
            _style_source_paragraph(date_line, 7.25, color="6B7C87")

            cells[2].text = f"« {_compact_excerpt(citation.text)} »"
            _style_source_paragraph(cells[2].paragraphs[0], 8.25, color="263238")

            cells[3].text = ""
            link_target = citation.deep_link or citation.url
            link_paragraph = cells[3].paragraphs[0]
            link_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_hyperlink(
                link_paragraph,
                link_target,
                "Ouvrir la source",
                font_size=8.25,
                bold=True,
            )
            domain = cells[3].add_paragraph(urlsplit(citation.url).netloc)
            domain.alignment = WD_ALIGN_PARAGRAPH.CENTER
            domain.paragraph_format.space_before = Pt(3)
            _style_source_paragraph(domain, 7, color="6B7C87")
    else:
        cells = table.add_row().cells
        cells[0].merge(cells[3]).text = "Aucune source contractuelle retenue."
        _style_source_paragraph(cells[0].paragraphs[0], 8.5, color="536878")

    _apply_table_geometry(table, _SOURCE_COLUMN_WIDTHS)
    _set_table_borders(table)


def render_grille(state: InterviewState) -> bytes:
    template = grille_template_path()
    if not template.is_file():
        raise RuntimeError(f"Gabarit Word introuvable : {template}")
    document = Document(template)
    _fill_identification(document, state)
    _fill_risk_table(table_by_heading(document, "Critère"), state, "A")
    if state.usages:
        for (heading, table), usage in zip(
            _usage_blocks(document, len(state.usages[:4])),
            state.usages[:4],
            strict=True,
        ):
            _fill_usage_heading(heading, usage)
            _fill_risk_table(table, state, "B")
    _append_source_register(document, state)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def write_grille(
    state: InterviewState, output_dir: str | os.PathLike | None = None
) -> Path:
    directory = Path(output_dir) if output_dir is not None else docx_output_dir()
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / grille_filename(state)
    path.write_bytes(render_grille(state))
    return path
