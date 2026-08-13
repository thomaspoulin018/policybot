from __future__ import annotations

import unicodedata
from collections.abc import Iterable
from urllib.parse import urlsplit

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor


def _normalized_text(value: str) -> str:
    folded = unicodedata.normalize("NFKD", value.replace("\xa0", " "))
    folded = "".join(char for char in folded if not unicodedata.combining(char))
    for dash in ("—", "–", "‑"):
        folded = folded.replace(dash, "-")
    for quote in ("’", "‘", "`"):
        folded = folded.replace(quote, "'")
    return " ".join(folded.casefold().split())


def table_by_heading(document, heading: str):
    """Return the table whose first cell starts with *heading*."""
    wanted = _normalized_text(heading)
    for table in document.tables:
        if table.rows and table.rows[0].cells:
            first_cell = _normalized_text(table.cell(0, 0).text)
            if first_cell.startswith(wanted):
                return table
    raise RuntimeError(
        "Le gabarit Word de grille d'évaluation ne contient aucun tableau "
        f"commençant par « {heading} »."
    )


def replace_paragraph_text(paragraph, old: str, new: str) -> None:
    """Replace text across runs while retaining the formatting of the first run."""
    current = paragraph.text
    if old not in current:
        raise RuntimeError(f"Texte introuvable dans le paragraphe : {old!r}")
    replacement = current.replace(old, new)
    if not paragraph.runs:
        paragraph.add_run(replacement)
        return
    paragraph.runs[0].text = replacement
    for run in paragraph.runs[1:]:
        run.text = ""


def tick_checkbox(paragraphs: Iterable, letter: str) -> None:
    wanted = letter.strip().casefold()
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text.casefold().endswith(wanted) and "☐" in text:
            replace_paragraph_text(paragraph, "☐", "☒")
            return
    raise RuntimeError(f"Case de risque introuvable pour le niveau « {letter} ».")


def add_hyperlink(
    paragraph,
    url: str,
    text: str | None = None,
    *,
    font_size: float = 8,
    bold: bool = False,
):
    """Append a styled external hyperlink to an existing paragraph."""
    if not url:
        return None
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), str(RGBColor(5, 99, 193)))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(font_size * 2)))
    size_complex = OxmlElement("w:szCs")
    size_complex.set(qn("w:val"), str(int(font_size * 2)))
    properties.extend((color, underline, size, size_complex))
    if bold:
        properties.append(OxmlElement("w:b"))
    run.append(properties)
    text_element = OxmlElement("w:t")
    parsed = urlsplit(url)
    text_element.text = text or (parsed.netloc + parsed.path).rstrip("/") or url
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_citation_paragraph(cell, label: str, quote: str, url: str):
    """Append a sourced quote and a clickable external link to a table cell."""
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    quote_run = paragraph.add_run(f"{label} — « {quote.strip()} »")
    quote_run.font.size = Pt(8)
    if not url:
        return paragraph

    paragraph.add_run().add_break()
    add_hyperlink(paragraph, url)
    return paragraph
